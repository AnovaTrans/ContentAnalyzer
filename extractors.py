import os
from typing import List, Dict
import pypdf
from docx import Document
from pptx import Presentation
import openpyxl
from lxml import etree
from langdetect import detect

class FileExtractor:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.filename = os.path.basename(file_path)
        self.extension = os.path.splitext(file_path)[1].lower()

    def extract(self) -> Dict:
        text = ""
        doc_type = "unknown"
        
        # Default Structure Metrics
        structure_metrics = {
            "table_count": 0,
            "image_count": 0,
            "chart_count": 0,
            "hyperlink_count": 0,
            "page_count": "Unknown",
            "ocr_needed": False
        }

        try:
            if self.extension in ['.docx', '.doc']:
                text, structure_metrics = self._extract_docx()
                doc_type = "office_word"
            elif self.extension == '.pdf':
                text, structure_metrics = self._extract_pdf()
                doc_type = "pdf"
            elif self.extension in ['.xlsx', '.xls']:
                text = self._extract_xlsx()
                doc_type = "office_excel"
            else:
                text = self._extract_text_based()
                doc_type = "text_plain"

            clean_text = text.strip()
            lang = self._detect_language(clean_text) if clean_text else "unknown"
            
            return {
                "filename": self.filename,
                "full_text": clean_text,
                "doc_type": doc_type,
                "detected_lang": lang,
                "structure_metrics": structure_metrics
            }

        except Exception as e:
            print(f"Error extracting {self.filename}: {str(e)}")
            return {
                "filename": self.filename,
                "full_text": "",
                "error": str(e),
                "structure_metrics": structure_metrics
            }

    def _detect_language(self, text: str) -> str:
        try: return detect(text[:1000])
        except: return "unknown"

    def _extract_docx(self):
        doc = Document(self.file_path)
        text = []
        
        # 1. Advanced Counting
        table_count = len(doc.tables)
        
        # Images (Inline & Floating)
        inline_shapes = doc.element.xpath('.//wp:inline')
        anchored_shapes = doc.element.xpath('.//wp:anchor')
        image_count = len(inline_shapes) + len(anchored_shapes)
        
        # Hyperlinks
        hyperlinks = doc.element.xpath('.//w:hyperlink')
        hyperlink_count = len(hyperlinks)

        # Text Extraction
        def extract_from_element(element):
            if hasattr(element, 'paragraphs'):
                for p in element.paragraphs:
                    if p.text.strip(): text.append(p.text)
            if hasattr(element, 'tables'):
                for table in element.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            extract_from_element(cell)

        extract_from_element(doc)
        
        meta = {
            "table_count": table_count,
            "image_count": image_count,
            "chart_count": 0, 
            "hyperlink_count": hyperlink_count,
            "page_count": "N/A (Flow)",
            "ocr_needed": image_count > 0 
        }
        return "\n".join(text), meta

    def _extract_pdf(self):
        reader = pypdf.PdfReader(self.file_path)
        text = []
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted: text.append(extracted)
            
        meta = {
            "table_count": "Unknown (PDF)",
            "image_count": len(reader.pages), 
            "chart_count": 0,
            "hyperlink_count": 0,
            "page_count": len(reader.pages),
            "ocr_needed": True 
        }
        return "\n".join(text), meta

    def _extract_xlsx(self) -> str:
        wb = openpyxl.load_workbook(self.file_path, data_only=True)
        text = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            for row in ws.iter_rows(values_only=True):
                row_text = [str(cell) for cell in row if cell is not None]
                if row_text: text.append(" ".join(row_text))
        return "\n".join(text)

    def _extract_text_based(self) -> str:
        with open(self.file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def _extract_pptx(self) -> str:
        return ""

# --- Missing Function Restored ---
def chunk_text(text: str, max_chars: int = 80000, overlap: int = 2000) -> List[str]:
    """Chunks text for LLM context windows. 
    Increased to 80K chars to reduce API calls while staying within model limits."""
    if not text:
        return []
    
    chunks = []
    start = 0
    while start < len(text):
        end = start + max_chars
        chunks.append(text[start:end])
        start += (max_chars - overlap)
    return chunks