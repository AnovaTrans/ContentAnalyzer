import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from models import DocumentAnalysis

class ReportGenerator:
    def __init__(self, analysis: DocumentAnalysis, output_dir: str):
        self.analysis = analysis
        self.output_dir = output_dir
        # Base name without extension
        raw_name = self.analysis.document_names[0]
        self.base_name = os.path.splitext(raw_name)[0].replace('.', '_').replace(' ', '_')
        # Timestamp format: YYYYMMDD_HHMMSS
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

    def generate_json(self):
        file_path = os.path.join(self.output_dir, f"{self.base_name}_{self.timestamp}_analysis.json")
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(self.analysis.model_dump_json(indent=4))
        print(f"Generated: {file_path}")

    def generate_docx(self):
        file_path = os.path.join(self.output_dir, f"{self.base_name}_{self.timestamp}_analysis_report.docx")
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Header
        title = doc.add_heading(f"Enterprise Document Analysis", 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(f"File: {self.base_name} | Date: {self.analysis.analysis_date}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph("---")

        def add_kv(label, value):
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(str(value))

        def add_list_simple(items):
            if not items: return
            for item in items: doc.add_paragraph(str(item), style='List Bullet')

        # 1. Summary & Corporate
        doc.add_heading("1. Executive Context", level=1)
        doc.add_paragraph(self.analysis.high_level_description)
        corp = self.analysis.corporate_context
        if corp.parent_company_or_group:
            p = doc.add_paragraph()
            p.add_run("🏢 Hierarchy: ").bold = True
            p.add_run(f"{corp.primary_company} ➤ {corp.parent_company_or_group}")
            
            p_brand = doc.add_paragraph()
            run_brand = p_brand.add_run(f"⚠️ Brand: {corp.brand_consistency_notes}")
            run_brand.font.color.rgb = RGBColor(200, 0, 0)

        # 2. Component Analysis
        doc.add_heading("2. Document Components & Structure", level=1)
        comp = self.analysis.component_analysis
        add_kv("Text Volume", comp.text_volume_summary)
        add_kv("Visuals", comp.visual_elements_summary)
        add_kv("Tables", comp.data_tables_summary)
        add_kv("OCR Needs", comp.ocr_requirements)
        add_kv("Formatting Impact", f"{comp.formatting_impact_score}/10")

        # 3. Domain Breakdown
        doc.add_heading("3. Domain Composition", level=1)
        for share in self.analysis.domain_breakdown:
            p = doc.add_paragraph()
            p.add_run(f"{share.percentage}% {share.domain}").bold = True
            p.add_run(f" - {share.justification}")

        # 4. Financials (Green Box)
        if self.analysis.financial_metrics:
            doc.add_heading("Commercial Estimation", level=1)
            fin = self.analysis.financial_metrics
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            cell = table.cell(0, 0)
            
            p = cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            runner = p.add_run(f"PRICE FACTOR: {fin.suggested_price_factor}x")
            runner.bold = True
            runner.font.size = Pt(16)
            runner.font.color.rgb = RGBColor(0, 100, 0)
            
            p_expl = cell.add_paragraph(f"Base Calculation: {fin.explanation}")
            p_expl.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            if fin.component_impact_breakdown:
                p_drivers = cell.add_paragraph()
                p_drivers.add_run("Cost Drivers:").bold = True
                for item in fin.component_impact_breakdown:
                    cell.add_paragraph(f"• {item}")

        # 5. Measurements & Protocols (NEW FIXED SECTION)
        doc.add_heading("4. Technical Protocols", level=1)
        meas = self.analysis.measurement_rules
        add_kv("Decimal Format", meas.decimal_format_rule)
        add_kv("Unit Conversion", meas.unit_conversion_rule)
        
        doc.add_heading("Critical Numbers Identified", level=3)
        add_list_simple(meas.critical_measurements_list)

        # 6. Localization Strategy (NEW SECTION)
        doc.add_heading("5. Localization Strategy", level=1)
        loc = self.analysis.localization_strategy
        add_kv("Cultural Adaptation", loc.cultural_adaptation_notes)
        add_kv("Geographic Handling", loc.geographic_handling)
        add_kv("Tone Guide", loc.tone_and_style_guide)

        # 7. Terminology (EXPANDED)
        doc.add_heading("6. Terminology", level=1)
        terms = self.analysis.terminology_and_resources.categorized_terms
        for category, t_list in terms.items():
            p = doc.add_paragraph()
            p.add_run(f"{category}: ").bold = True
            p.add_run(", ".join(t_list))

        # 8. Resource Requirements
        doc.add_heading("7. Resource Qualifications", level=1)
        res = self.analysis.recommended_resources
        add_kv("Profile", res.headline_profile)
        add_kv("Seniority", res.seniority_level)
        add_kv("Min Experience", f"{res.min_years_experience} Years")
        add_kv("Education", res.education_requirement)
        add_kv("Tools", ", ".join(res.tool_recommendations))

        # 9. Risks
        doc.add_heading("8. Risk Matrix", level=1)
        for risk in self.analysis.risks_and_mitigations:
            p = doc.add_paragraph()
            p.add_run(f"[{risk.category}] {risk.name}: ").bold = True
            p.add_run(risk.description)
            p_ev = doc.add_paragraph(style='List Bullet')
            p_ev.add_run(f"Evidence: \"{risk.evidence_from_text}\"").italic = True

        # 10. Workflow
        doc.add_heading("9. Workflow", level=1)
        wf = self.analysis.workflow_recommendations
        add_kv("Pre-Processing", ", ".join(wf.pre_processing))
        add_kv("Translation", str(wf.translation_process))
        add_kv("QA", ", ".join(wf.qa_and_review))
        add_kv("Post-Processing", ", ".join(wf.post_processing))
        
        doc.save(file_path)
        print(f"Generated Enterprise DOCX: {file_path}")